# -*- coding: utf-8 -*-
"""
Word文档文本提取器
"""

from pathlib import Path
import subprocess
import tempfile
import os

from .base import BaseExtractor, ExtractedDocument, ExtractedPage
from ..utils.logger import get_logger


class DocExtractor(BaseExtractor):
    """Word文档文本提取器"""

    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.doc', '.docx', '.rtf', '.wps']
        self.logger = get_logger("doc_extractor")

    def extract(self, file_path: Path) -> ExtractedDocument:
        """提取Word文档内容"""
        file_path = Path(file_path)

        if not file_path.exists():
            return self._create_error_document(file_path, f"文件不存在: {file_path}")

        suffix = file_path.suffix.lower()

        try:
            if suffix == '.docx':
                return self._extract_docx(file_path)
            elif suffix in ['.doc', '.rtf', '.wps']:
                return self._extract_doc(file_path)
            else:
                return self._create_error_document(
                    file_path, f"不支持的文件格式: {suffix}"
                )
        except Exception as e:
            self.logger.error(f"提取Word文档失败 {file_path}: {e}")
            return self._create_error_document(file_path, str(e))

    def _extract_docx(self, file_path: Path) -> ExtractedDocument:
        """提取docx文件"""
        from docx import Document

        doc = Document(file_path)
        paragraphs = []

        # 提取段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # 提取表格
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    paragraphs.append(" | ".join(row_text))

        full_text = '\n\n'.join(paragraphs)
        full_text = self._clean_text(full_text)

        # docx没有明确的分页概念，作为单页处理
        pages = [ExtractedPage(
            page_num=1,
            content=full_text,
            has_images=self._has_images_docx(doc),
            has_tables=len(doc.tables) > 0,
        )]

        return ExtractedDocument(
            source_path=file_path,
            file_name=file_path.name,
            file_type='.docx',
            pages=pages,
            full_text=full_text,
            metadata={'engine': 'python-docx'},
        )

    def _extract_doc(self, file_path: Path) -> ExtractedDocument:
        """提取doc文件（旧格式）"""
        # 尝试多种方法
        text = None
        engine_used = None
        errors = []

        # 方法1: 尝试使用 antiword (Linux/Mac)
        try:
            text = self._extract_with_antiword(file_path)
            engine_used = 'antiword'
        except Exception as e:
            errors.append(f"antiword失败: {e}")

        # 方法2: 尝试使用 textract
        if text is None:
            try:
                text = self._extract_with_textract(file_path)
                engine_used = 'textract'
            except Exception as e:
                errors.append(f"textract失败: {e}")

        # 方法3: 尝试使用 win32com (Windows)
        if text is None:
            try:
                text = self._extract_with_win32com(file_path)
                engine_used = 'win32com'
            except Exception as e:
                errors.append(f"win32com失败: {e}")

        # 方法4: 尝试直接读取（可能有乱码）
        if text is None:
            try:
                text = self._extract_raw(file_path)
                engine_used = 'raw'
            except Exception as e:
                errors.append(f"raw读取失败: {e}")

        if text is None:
            return ExtractedDocument(
                source_path=file_path,
                file_name=file_path.name,
                file_type=file_path.suffix.lower(),
                pages=[],
                full_text="",
                errors=errors,
            )

        text = self._clean_text(text)

        pages = [ExtractedPage(
            page_num=1,
            content=text,
            has_images=False,
            has_tables=False,
        )]

        return ExtractedDocument(
            source_path=file_path,
            file_name=file_path.name,
            file_type=file_path.suffix.lower(),
            pages=pages,
            full_text=text,
            metadata={'engine': engine_used},
            errors=errors if engine_used == 'raw' else [],
        )

    def _extract_with_antiword(self, file_path: Path) -> str:
        """使用antiword提取"""
        result = subprocess.run(
            ['antiword', str(file_path)],
            capture_output=True,
            text=True,
            timeout=30
        )
        if result.returncode != 0:
            raise Exception(result.stderr)
        return result.stdout

    def _extract_with_textract(self, file_path: Path) -> str:
        """使用textract提取"""
        import textract
        text = textract.process(str(file_path))
        return text.decode('utf-8')

    def _extract_with_win32com(self, file_path: Path) -> str:
        """使用win32com提取（Windows专用）"""
        import win32com.client
        import pythoncom

        pythoncom.CoInitialize()
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False

            doc = word.Documents.Open(str(file_path.absolute()))
            text = doc.Content.Text
            doc.Close()
            word.Quit()

            return text
        finally:
            pythoncom.CoUninitialize()

    def _extract_raw(self, file_path: Path) -> str:
        """直接读取文件（尝试提取可读文本）"""
        with open(file_path, 'rb') as f:
            content = f.read()

        # 尝试提取ASCII/Unicode文本
        import re
        # 匹配可打印的中英文字符序列
        text_parts = re.findall(
            rb'[\x20-\x7e\u4e00-\u9fff]+',
            content
        )

        texts = []
        for part in text_parts:
            try:
                decoded = part.decode('utf-8')
                if len(decoded) > 5:  # 过滤太短的片段
                    texts.append(decoded)
            except:
                try:
                    decoded = part.decode('gbk')
                    if len(decoded) > 5:
                        texts.append(decoded)
                except:
                    pass

        return '\n'.join(texts)

    def _has_images_docx(self, doc) -> bool:
        """检查docx是否有图片"""
        try:
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    return True
        except:
            pass
        return False
